"""Document processing service - PDF & DOCX parsing with chunking."""

import re
from pathlib import Path
from typing import List, Dict, Optional
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(pdf_path: Path) -> str:
    """
    Extract text from PDF file.
    
    Args:
        pdf_path: Path to PDF file
        
    Returns:
        Extracted text content
    """
    try:
        import PyPDF2
        
        text = []
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            
            logger.info(f"📄 Processing PDF: {total_pages} pages")
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text.append(page_text)
                    logger.debug(f"  Page {page_num}: {len(page_text)} chars")
        
        full_text = "\n\n".join(text)
        logger.info(f"✅ Extracted {len(full_text)} characters from PDF")
        return full_text
        
    except Exception as e:
        logger.error(f"❌ PDF extraction failed: {e}")
        raise RuntimeError(f"Failed to extract PDF: {e}")


def extract_text_from_docx(docx_path: Path) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        from docx import Document
        
        doc = Document(str(docx_path))
        text = []
        
        logger.info(f"📄 Processing DOCX: {len(doc.paragraphs)} paragraphs")
        
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
        
        full_text = "\n\n".join(text)
        logger.info(f"✅ Extracted {len(full_text)} characters from DOCX")
        return full_text
        
    except Exception as e:
        logger.error(f"❌ DOCX extraction failed: {e}")
        raise RuntimeError(f"Failed to extract DOCX: {e}")


def chunk_text(
    text: str,
    chunk_size: int = 1000,
    overlap: int = 200,
    min_chunk_size: int = 100
) -> List[Dict[str, any]]:
    """
    Chunk text into overlapping segments for embedding.
    
    Args:
        text: Text to chunk
        chunk_size: Target chunk size in characters
        overlap: Overlap between chunks in characters
        min_chunk_size: Minimum chunk size to keep
        
    Returns:
        List of chunk dicts with 'text', 'start', 'end', 'chunk_id'
    """
    # Clean and normalize text
    text = re.sub(r'\s+', ' ', text).strip()
    
    chunks = []
    start = 0
    chunk_id = 0
    
    while start < len(text):
        # Calculate end position
        end = start + chunk_size
        
        # If this is not the last chunk, try to break at sentence boundary
        if end < len(text):
            # Look for sentence endings within the chunk
            sentence_endings = ['.', '!', '?', '\n']
            best_break = end
            
            # Search backwards from end for a sentence boundary
            for i in range(end, max(start + min_chunk_size, end - 100), -1):
                if i < len(text) and text[i] in sentence_endings:
                    # Check if next char is space or end of text
                    if i + 1 >= len(text) or text[i + 1].isspace():
                        best_break = i + 1
                        break
            
            end = best_break
        
        # Extract chunk
        chunk_text = text[start:end].strip()
        
        if len(chunk_text) >= min_chunk_size:
            chunks.append({
                'chunk_id': chunk_id,
                'text': chunk_text,
                'start': start,
                'end': end,
                'length': len(chunk_text)
            })
            chunk_id += 1
        
        # Move start position with overlap
        start = end - overlap
        
        # Ensure we make progress
        if start >= end:
            start = end
    
    logger.info(f"✅ Created {len(chunks)} chunks from {len(text)} characters")
    return chunks


def extract_topics_from_text(text: str, max_topics: int = 20) -> List[str]:
    """
    Extract potential topics from text using simple heuristics.
    
    Args:
        text: Text to analyze
        max_topics: Maximum number of topics to extract
        
    Returns:
        List of topic strings
    """
    topics = []
    
    # Look for headers (lines with fewer than 80 chars, title case, no punctuation at end)
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if 10 < len(line) < 80 and line[0].isupper() and line[-1] not in '.!?:;,':
            # Check if it looks like a header (mostly capitalized words)
            words = line.split()
            if len(words) > 1 and sum(1 for w in words if w[0].isupper()) / len(words) > 0.5:
                topics.append(line)
    
    # If we found too few, extract noun phrases from first sentences
    if len(topics) < 5:
        # Simple extraction: capitalized phrases
        import re
        capitalized_phrases = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b', text[:5000])
        topics.extend(list(set(capitalized_phrases))[:max_topics - len(topics)])
    
    logger.info(f"📋 Extracted {len(topics[:max_topics])} potential topics")
    return topics[:max_topics]


def process_document(file_path: Path) -> Dict:
    """
    Process a document (PDF/DOCX) and extract text, chunks, and topics.
    
    Args:
        file_path: Path to document file
        
    Returns:
        Dict with:
            - 'filename': Original filename
            - 'text': Full extracted text
            - 'chunks': List of text chunks
            - 'topics': List of potential topics
            - 'char_count': Total character count
    """
    suffix = file_path.suffix.lower()
    
    logger.info(f"🔄 Processing document: {file_path.name}")
    
    # Extract text based on file type
    if suffix == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif suffix == '.docx':
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
    
    # Chunk the text
    chunks = chunk_text(text, chunk_size=1000, overlap=200)
    
    # Extract topics
    topics = extract_topics_from_text(text)
    
    result = {
        'filename': file_path.name,
        'file_type': suffix[1:],  # Remove dot
        'text': text,
        'chunks': chunks,
        'topics': topics,
        'char_count': len(text),
        'chunk_count': len(chunks)
    }
    
    logger.info(f"✅ Document processed: {result['char_count']} chars, "
                f"{result['chunk_count']} chunks, {len(topics)} topics")
    
    return result
